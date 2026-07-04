import asyncio
import hashlib
import os
from pathlib import Path
from typing import Iterable, List

from dotenv import load_dotenv
from docx import Document as DocxDocument
from pypdf import PdfReader

from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter


load_dotenv()


PROJECT_ROOT = Path(__file__).resolve().parents[3]

DEFAULT_DOCS_DIR = PROJECT_ROOT / "data" / "rag_documents"

PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME")
PINECONE_NAMESPACE = os.getenv("PINECONE_NAMESPACE", "activiti-mediation-template-docs")

if not PINECONE_API_KEY:
    raise RuntimeError("PINECONE_API_KEY is not configured in .env")

if not PINECONE_INDEX_NAME:
    raise RuntimeError("PINECONE_INDEX_NAME is not configured in .env")


embeddings = OpenAIEmbeddings(
    model="text-embedding-3-small",
    show_progress_bar=False,
    chunk_size=50,
    retry_min_seconds=10,
)

vectorstore = PineconeVectorStore(
    index_name=PINECONE_INDEX_NAME,
    embedding=embeddings,
    namespace=PINECONE_NAMESPACE,
)


def _stable_id(text: str, metadata: dict) -> str:
    """
    Creates a stable chunk ID so re-ingestion updates the same vectors
    instead of creating uncontrolled duplicates.
    """
    source = str(metadata.get("source", "unknown"))
    chunk_index = str(metadata.get("chunk_index", "0"))
    raw = f"{source}|{chunk_index}|{text}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _detect_doc_category(file_path: Path) -> str:
    """
    Helps retrieval later by tagging what kind of document this is.
    """
    name = file_path.name.lower()

    if "activiti" in name and "expression" in name:
        return "expression_guide"

    if "objective" in name:
        return "objective"

    if "template" in name:
        return "template_data"

    return "general"


def load_docx(file_path: Path) -> Document:
    """
    Loads a .docx file into one LangChain Document.
    """
    doc = DocxDocument(file_path)

    paragraphs = [
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.text and paragraph.text.strip()
    ]

    text = "\n".join(paragraphs)

    return Document(
        page_content=text,
        metadata={
            "source": str(file_path),
            "source_file": file_path.name,
            "file_type": "docx",
            "doc_category": _detect_doc_category(file_path),
        },
    )


def load_pdf(file_path: Path) -> List[Document]:
    """
    Loads a .pdf file into one Document per page.
    """
    reader = PdfReader(str(file_path))
    documents: List[Document] = []

    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""

        if not text.strip():
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": str(file_path),
                    "source_file": file_path.name,
                    "file_type": "pdf",
                    "page_number": page_index,
                    "doc_category": _detect_doc_category(file_path),
                },
            )
        )

    return documents


def load_text_file(file_path: Path) -> Document:
    """
    Loads .txt, .md, and .sql files.
    """
    text = file_path.read_text(encoding="utf-8", errors="ignore")

    return Document(
        page_content=text,
        metadata={
            "source": str(file_path),
            "source_file": file_path.name,
            "file_type": file_path.suffix.replace(".", "").lower(),
            "doc_category": _detect_doc_category(file_path),
        },
    )


def load_single_document(file_path: Path) -> List[Document]:
    """
    Loads one supported document file.
    """
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return [load_docx(file_path)]

    if suffix == ".pdf":
        return load_pdf(file_path)

    if suffix in {".txt", ".md", ".sql"}:
        return [load_text_file(file_path)]

    print(f"Skipping unsupported file: {file_path}")
    return []


def iter_supported_files(folder_path: Path) -> Iterable[Path]:
    supported_extensions = {".docx", ".pdf", ".txt", ".md", ".sql"}

    for file_path in folder_path.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            yield file_path


def load_documents_from_folder(folder_path: Path) -> List[Document]:
    """
    Loads all supported documents from a folder.
    """
    all_documents: List[Document] = []

    for file_path in iter_supported_files(folder_path):
        print(f"Loading document: {file_path.name}")
        loaded_documents = load_single_document(file_path)
        all_documents.extend(loaded_documents)

    return all_documents


def split_documents(documents: List[Document]) -> List[Document]:
    """
    Splits loaded documents into chunks.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2500,
        chunk_overlap=300,
        separators=[
            "\n## ",
            "\n### ",
            "\n\n",
            "\n",
            ". ",
            " ",
            "",
        ],
    )

    chunks = text_splitter.split_documents(documents)

    for index, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = index
        chunk.metadata["chunk_id"] = _stable_id(
            text=chunk.page_content,
            metadata=chunk.metadata,
        )

    return chunks

def delete_existing_vectors_for_documents(documents: List[Document]) -> None:
    """
    Deletes existing Pinecone vectors for the same source files before re-ingestion.

    This avoids duplicate/stale chunks when the same document is ingested again.
    """
    source_files = sorted(
        {
            str(doc.metadata.get("source_file"))
            for doc in documents
            if doc.metadata.get("source_file")
        }
    )

    if not source_files:
        print("No source_file metadata found. Skipping Pinecone cleanup.")
        return

    print("PINECONE CLEANUP PHASE")

    for source_file in source_files:
        print(f"Deleting existing vectors for source_file: {source_file}")

        vectorstore.delete(
            filter={
                "source_file": {
                    "$eq": source_file,
                }
            }
        )


async def index_documents_async(
    documents: List[Document],
    batch_size: int = 50,
) -> None:
    """
    Adds document chunks to Pinecone in async batches.
    """
    print("VECTOR STORAGE PHASE")
    print(
        f"VectorStore Indexing: Preparing to add {len(documents)} chunks "
        f"to Pinecone index '{PINECONE_INDEX_NAME}' namespace '{PINECONE_NAMESPACE}'"
    )

    batches = [
        documents[i : i + batch_size]
        for i in range(0, len(documents), batch_size)
    ]

    print(
        f"VectorStore Indexing: Split into {len(batches)} batches "
        f"of {batch_size} chunks each"
    )

    async def add_batch(batch: List[Document], batch_num: int) -> bool:
        try:
            ids = [doc.metadata["chunk_id"] for doc in batch]
            await vectorstore.aadd_documents(batch, ids=ids)

            print(
                f"VectorStore Indexing: Successfully added batch "
                f"{batch_num}/{len(batches)} ({len(batch)} chunks)"
            )
            return True

        except Exception as exc:
            print(f"VectorStore Indexing: Failed to add batch {batch_num} - {exc}")
            return False

    tasks = [
        add_batch(batch, i + 1)
        for i, batch in enumerate(batches)
    ]

    results = await asyncio.gather(*tasks, return_exceptions=True)

    successful = sum(1 for result in results if result is True)

    if successful == len(batches):
        print(
            f"VectorStore Indexing: All batches processed successfully! "
            f"({successful}/{len(batches)})"
        )
    else:
        print(
            f"VectorStore Indexing: Processed {successful}/{len(batches)} "
            f"batches successfully"
        )


async def main() -> None:
    """
    Main ingestion function.

    Default input folder:
    data/rag_documents
    """
    print("ACTIVITI MEDIATION DOCUMENT INGESTION PIPELINE")

    docs_dir = Path(os.getenv("RAG_DOCUMENTS_DIR", str(DEFAULT_DOCS_DIR))).resolve()

    if not docs_dir.exists():
        raise RuntimeError(f"RAG documents folder does not exist: {docs_dir}")

    print(f"Loading documents from: {docs_dir}")

    all_docs = load_documents_from_folder(docs_dir)

    if not all_docs:
        print("No supported documents found. Nothing to ingest.")
        return
    
    delete_existing_vectors_for_documents(all_docs)

    chunks = split_documents(all_docs)

    print(f"Text Splitter: Created {len(chunks)} chunks from {len(all_docs)} documents")

    await index_documents_async(chunks, batch_size=50)

    print("PIPELINE COMPLETE")
    print("Documentation ingestion pipeline finished successfully!")
    print("Summary:")
    print(f"   • Documents loaded: {len(all_docs)}")
    print(f"   • Chunks created: {len(chunks)}")
    print(f"   • Pinecone index: {PINECONE_INDEX_NAME}")
    print(f"   • Pinecone namespace: {PINECONE_NAMESPACE}")


if __name__ == "__main__":
    asyncio.run(main())